# =============================================================================
# agents/policy_drafter/docx_builder.py
# Converts AI-generated draft sections into a properly formatted .docx file.
# Called after draft_document() produces the sections dict.
# Returns a BytesIO buffer ready for HTTP response / SharePoint upload.
# =============================================================================

import io
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
#  Colour palette — Dragnet brand
# =============================================================================

BRAND_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black navy
BRAND_ACCENT = RGBColor(0x37, 0x8A, 0xDD)   # Dragnet blue
BRAND_MID    = RGBColor(0x44, 0x47, 0x5A)   # dark grey
BRAND_LIGHT  = RGBColor(0xF4, 0xF6, 0xFA)   # off-white background
RULE_GREY    = RGBColor(0xCC, 0xCC, 0xCC)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
RED_DRAFT    = RGBColor(0xA3, 0x2D, 0x2D)


# =============================================================================
#  Helpers
# =============================================================================

def _set_para_border_bottom(para, color: str = "CCCCCC", size: int = 4):
    """Add a bottom border rule to a paragraph (replaces table-as-divider)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_shading(cell, fill_hex: str):
    """Set cell background shading (ShadingType.CLEAR pattern)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "CCCCCC"):
    """Set thin borders on a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _run(para, text: str, bold=False, italic=False,
         size_pt: int = 10, color: RGBColor = None, font: str = "Arial"):
    run = para.add_run(text)
    run.bold  = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _heading(doc: Document, text: str, level: int = 1):
    """
    Add a CDI-style section heading with a bottom rule.
    level 1 = major section (blue, 11pt bold)
    level 2 = sub-section (dark, 10pt bold)
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(2)
    if level == 1:
        _run(para, text.upper(), bold=True, size_pt=11, color=BRAND_ACCENT)
        _set_para_border_bottom(para, color="378ADD", size=6)
    else:
        _run(para, text, bold=True, size_pt=10, color=BRAND_DARK)
    return para


def _body(doc: Document, text: str, indent: bool = False):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    _run(para, text, size_pt=10, color=BRAND_MID)
    return para


def _bullet(doc: Document, text: str, numbered: bool = False, num_val: int = 1):
    """Add a properly formatted bullet or numbered list item."""
    para = doc.add_paragraph(style="List Bullet" if not numbered else "List Number")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Inches(0.3)
    _run(para, text.lstrip("•-– 0123456789."), size_pt=10, color=BRAND_MID)
    return para


def _parse_and_add_lines(doc: Document, text: str, is_policy_statement: bool = False):
    """
    Intelligently render multi-line AI output:
    - Numbered lines (1. / 2.) → numbered list
    - Dash/bullet lines (- / • / *) → bullet list
    - Sub-role headers (e.g. "Compliance Lead\n-") → bold role + bullets
    - Plain paragraphs → body text
    """
    lines = [l.rstrip() for l in text.splitlines()]
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Numbered item: "1." "2." etc.
        if len(line) > 2 and line[0].isdigit() and line[1] in ".)" :
            content = line[2:].strip() if len(line) > 2 else line
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.left_indent  = Inches(0.3)
            _run(para, content, size_pt=10, color=BRAND_MID)

        # Bullet item: "- " or "• " or "* "
        elif line.startswith(("-", "•", "*")) and len(line) > 2:
            content = line.lstrip("-•* ").strip()
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.left_indent  = Inches(0.3)
            _run(para, content, size_pt=10, color=BRAND_MID)

        # Role header in responsibilities section (not starting with dash/number/blank)
        # Heuristic: all-caps or title-case short line followed by bullet lines
        elif (
            not line.startswith(("-", "•", "*"))
            and not (line[0].isdigit() and line[1:2] in (".", ")"))
            and len(line) < 60
            and i + 1 < len(lines)
            and lines[i + 1].strip().startswith("-")
        ):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after  = Pt(2)
            _run(para, line, bold=True, size_pt=10, color=BRAND_DARK)

        # Plain body paragraph
        else:
            _body(doc, line)

        i += 1


# =============================================================================
#  Cover page
# =============================================================================

def _add_cover(doc: Document, draft: dict):
    """Add a styled cover page."""
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, "DRAGNET SOLUTIONS LIMITED", bold=True, size_pt=14, color=BRAND_DARK)

    # Horizontal rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_border_bottom(rule, color="378ADD", size=8)

    # Document title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    p2.paragraph_format.space_after  = Pt(8)
    _run(p2, draft["title"].upper(), bold=True, size_pt=16, color=BRAND_ACCENT)

    # DRAFT watermark badge
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(32)
    _run(p3, "  DRAFT — AI GENERATED — PENDING REVIEW  ",
         bold=True, size_pt=9, color=RED_DRAFT)

    # Metadata table
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = "Table Grid"

    col_w = [Inches(2.0), Inches(4.5)]
    for row in meta_table.rows:
        row.cells[0].width = col_w[0]
        row.cells[1].width = col_w[1]

    rows_data = [
        ("Document Code",  draft["doc_code"]),
        ("Document Type",  draft["doc_type"]),
        ("Department",     draft["department"]),
        ("Version",        "1.0"),
        ("Status",         "DRAFT"),
        ("Date",           "[DATE — to be completed by owner]"),
    ]

    for i, (label, value) in enumerate(rows_data):
        c0, c1 = meta_table.rows[i].cells
        _set_cell_shading(c0, "1A1A2E")
        _set_cell_shading(c1, "F4F6FA")
        _set_cell_borders(c0)
        _set_cell_borders(c1)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(4)
        _run(p0, label, bold=True, size_pt=9, color=WHITE)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after  = Pt(4)
        _run(p1, value, size_pt=10,
             color=BRAND_DARK if label != "Status" else RED_DRAFT,
             bold=(label == "Status"))

    # Standards
    if draft.get("standards_mapping"):
        doc.add_paragraph()
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p4, f"Standards: {draft['standards_mapping']}", size_pt=9,
             color=BRAND_MID, italic=True)

    doc.add_page_break()


# =============================================================================
#  Revision history table
# =============================================================================

def _add_revision_history(doc: Document, draft: dict):
    _heading(doc, "Revision History", level=1)

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Inches(0.7), Inches(1.5), Inches(1.8), Inches(2.5)]
    headers = ["Version", "Date", "Author", "Change"]

    # Header row
    hdr_row = tbl.rows[0]
    for j, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[j]
        cell.width = w
        _set_cell_shading(cell, "1A1A2E")
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _run(p, hdr, bold=True, size_pt=9, color=WHITE)

    # Data row
    data_row = tbl.rows[1]
    data = ["1.0", "[DATE]", "[AUTHOR]", "Initial draft — AI-generated"]
    for j, (val, w) in enumerate(zip(data, col_widths)):
        cell = data_row.cells[j]
        cell.width = w
        _set_cell_shading(cell, "F4F6FA")
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _run(p, val, size_pt=9, color=BRAND_MID)

    doc.add_paragraph()


# =============================================================================
#  Review and approval table
# =============================================================================

def _add_review_approval(doc: Document):
    _heading(doc, "8. Review and Approval", level=1)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Inches(2.5), Inches(4.0)]
    rows_data = [
        ("Document Owner",    "[ROLE FROM ROLE REGISTER]"),
        ("Approved By",       "[APPROVER NAME AND ROLE]"),
        ("Effective Date",    "[DATE]"),
        ("Next Review Date",  "[DATE + 12 MONTHS]"),
        ("Classification",    "Internal"),
    ]

    for i, (label, value) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        _set_cell_shading(c0, "F4F6FA")
        _set_cell_shading(c1, "FFFFFF")
        _set_cell_borders(c0)
        _set_cell_borders(c1)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(4)
        _run(p0, label, bold=True, size_pt=9, color=BRAND_DARK)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after  = Pt(4)
        _run(p1, value, size_pt=9, color=BRAND_MID, italic=True)


# =============================================================================
#  Header / Footer
# =============================================================================

def _add_header_footer(doc: Document, draft: dict):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(hp, f"{draft['doc_code']}  |  v1.0 DRAFT", size_pt=8,
         color=BRAND_MID, italic=True)
    _set_para_border_bottom(hp, color="CCCCCC", size=4)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, "DRAGNET SOLUTIONS LIMITED  |  CONFIDENTIAL — INTERNAL USE ONLY  |  ",
         size_pt=8, color=BRAND_MID)
    # Page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = BRAND_MID
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# =============================================================================
#  Main builder
# =============================================================================

def build_docx(draft: dict) -> io.BytesIO:
    """
    Build a formatted .docx from a draft dict produced by service.py.

    Parameters
    ----------
    draft : dict
        Must contain: doc_code, title, doc_type, department,
        standards_mapping, sections {purpose, scope, policy_statement,
        responsibilities, procedure, records}

    Returns
    -------
    io.BytesIO
        Ready to send as HTTP response or upload to SharePoint.
    """
    doc = Document()

    # Page setup — A4 with 1-inch margins
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin   = Inches(1)
        section.bottom_margin = Inches(0.8)

    # Default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Cover page ───────────────────────────────────────────────────────────
    _add_cover(doc, draft)

    # ── Header / Footer ──────────────────────────────────────────────────────
    _add_header_footer(doc, draft)

    # ── Revision history ─────────────────────────────────────────────────────
    _add_revision_history(doc, draft)

    sections = draft.get("sections", {})

    # ── 1. Purpose ───────────────────────────────────────────────────────────
    _heading(doc, "1. Purpose", level=1)
    _parse_and_add_lines(doc, sections.get("purpose", "[Purpose not generated]"))
    doc.add_paragraph()

    # ── 2. Scope ─────────────────────────────────────────────────────────────
    _heading(doc, "2. Scope", level=1)
    _parse_and_add_lines(doc, sections.get("scope", "[Scope not generated]"))
    doc.add_paragraph()

    # ── 3. Policy Statement ──────────────────────────────────────────────────
    _heading(doc, "3. Policy Statement", level=1)
    _parse_and_add_lines(doc, sections.get("policy_statement", "[Policy statement not generated]"),
                         is_policy_statement=True)
    doc.add_paragraph()

    # ── 4. Responsibilities ──────────────────────────────────────────────────
    _heading(doc, "4. Responsibilities", level=1)
    _parse_and_add_lines(doc, sections.get("responsibilities", "[Responsibilities not generated]"))
    doc.add_paragraph()

    # ── 5. Procedure ─────────────────────────────────────────────────────────
    _heading(doc, "5. Procedure", level=1)
    _parse_and_add_lines(doc, sections.get("procedure", "Refer to the associated procedure document."))
    doc.add_paragraph()

    # ── 6. Records ───────────────────────────────────────────────────────────
    _heading(doc, "6. Records", level=1)
    records_text = sections.get("records", "")
    if records_text:
        # Records are usually "Name (Type: X) — Source: Y — Retain: Z" lines
        for line in records_text.splitlines():
            line = line.strip()
            if not line:
                continue
            # Render each record as a bullet
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.left_indent  = Inches(0.3)
            _run(para, line.lstrip("-•* "), size_pt=9, color=BRAND_MID)
    else:
        _body(doc, "[Records not generated]")
    doc.add_paragraph()

    # ── 7. Related Documents ──────────────────────────────────────────────────
    _heading(doc, "7. Related Documents", level=1)
    _body(doc, "[To be completed by document owner]", indent=False)
    doc.add_paragraph()

    # ── 8. Review and Approval ────────────────────────────────────────────────
    _add_review_approval(doc)

    # ── End marker ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_border_bottom(end_para, color="378ADD", size=4)
    _run(end_para, f"END OF DOCUMENT — {draft['doc_code']} v1.0 DRAFT",
         size_pt=8, color=BRAND_MID, italic=True)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf